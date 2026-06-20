"""
Test Plans, Scenarios, and Test Cases endpoints.
"""
import json
import io
from docx import Document
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import JSONResponse, StreamingResponse
from sqlalchemy import select

from app.api.deps import CurrentUser, DBSession, require_entity_permission, require_entity_project_access, require_permission, require_project_access
from app.config import get_settings
from app.models.requirement import Requirement
from app.models.test_plan import TestPlan
from app.models.test_scenario import TestScenario
from app.models.test_case import TestCase
from app.models.project import Project
from app.schemas.test_plan import (
    TestPlanOut, TestPlanUpdate,
    TestScenarioOut,
    TestCaseOut, TestCaseUpdate, TestCaseHistoryOut, TestCaseJiraSyncOut, TestCaseSummaryOut,
    AgentPlanTrigger, AgentCaseTrigger,
)
from app.schemas.requirement import ApprovalRequest
from app.services import agent_run_service, approval_service, test_plan_service, traceability_service
from app.services.agent_dispatch_service import enqueue_agent_run
from app.services.display_id_service import display_id, temporary_id
from app.services.rbac_service import APPROVE_TEST_CASES, APPROVE_TEST_PLANS, SYNC_JIRA
from app.agents.test_planning.planning_agent import TestPlanningAgent
from app.agents.test_planning.scenario_agent import TestScenarioAgent
from app.agents.test_planning.test_case_agent import TestCaseDevelopmentAgent

router = APIRouter()
settings = get_settings()


def _run_agents_synchronously() -> bool:
    return False


# ── GAP-4c: Quality gate ──────────────────────────────────────────────────────
QUALITY_GATE_MIN_SCORE = 2.5


def _quality_gate_blockers(req_rows: list[Requirement]) -> list[dict]:
    """Return requirements that should not drive generation without an override.

    Blocked when: rejected, AI quality verdict 'fail', or quality score below
    threshold. Requirements without a quality review are allowed (soft gate).
    """
    blocked = []
    for r in req_rows:
        reasons = []
        if r.status == "rejected":
            reasons.append("requirement is rejected")
        if r.quality_verdict == "fail":
            reasons.append("AI quality verdict is 'fail'")
        elif r.quality_score is not None and r.quality_score < QUALITY_GATE_MIN_SCORE:
            reasons.append(f"quality score {r.quality_score} is below {QUALITY_GATE_MIN_SCORE}")
        if reasons:
            blocked.append({
                "id": r.id,
                "requirement_id": r.requirement_id,
                "title": r.title,
                "reasons": reasons,
            })
    return blocked


def _enforce_quality_gate(req_rows: list[Requirement], override: bool) -> None:
    blocked = _quality_gate_blockers(req_rows)
    if blocked and not override:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "quality_gate_blocked",
                "message": (
                    f"{len(blocked)} requirement(s) failed the quality gate. "
                    "Improve the requirement(s), re-run the quality review, or retry with override_quality_gate=true."
                ),
                "blocked_requirements": blocked,
            },
        )


# ── Test Plans ────────────────────────────────────────────────────────────────

@router.get("/project/{project_id}", response_model=list[TestPlanOut])
async def list_test_plans(
    project_id: int,
    db: DBSession,
    current_user: CurrentUser,
):
    await require_project_access(project_id, current_user, db)
    return await test_plan_service.list_test_plans(db, project_id)


@router.get("/{plan_id}", response_model=TestPlanOut)
async def get_test_plan(
    plan_id: int,
    db: DBSession,
    current_user: CurrentUser,
):
    plan = await test_plan_service.get_test_plan(db, plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="Test plan not found")
    await require_entity_permission(plan, APPROVE_TEST_PLANS, current_user, db)
    return plan


@router.patch("/{plan_id}", response_model=TestPlanOut)
async def update_test_plan(
    plan_id: int,
    updates: TestPlanUpdate,
    db: DBSession,
    current_user: CurrentUser,
):
    plan = await test_plan_service.get_test_plan(db, plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="Test plan not found")
    await require_entity_permission(plan, APPROVE_TEST_PLANS, current_user, db)
    plan = await test_plan_service.update_test_plan(db, plan, updates)
    await db.commit()
    return plan


@router.post("/{plan_id}/approve", response_model=TestPlanOut)
async def approve_test_plan(
    plan_id: int,
    body: ApprovalRequest,
    db: DBSession,
    current_user: CurrentUser,
):
    plan = await test_plan_service.get_test_plan(db, plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="Test plan not found")
    await require_entity_permission(plan, APPROVE_TEST_PLANS, current_user, db)
    plan = await test_plan_service.approve_test_plan(db, plan, body.action, body.notes)
    await approval_service.create_approval_action(
        db,
        project_id=plan.project_id,
        user_id=current_user.id,
        entity_type="test_plan",
        entity_id=plan.id,
        action=body.action,
        notes=body.notes,
    )
    await db.commit()
    return plan


# ── Scenarios ─────────────────────────────────────────────────────────────────

@router.get("/scenarios/project/{project_id}", response_model=list[TestScenarioOut])
async def list_scenarios(
    project_id: int,
    db: DBSession,
    current_user: CurrentUser,
    requirement_id: int | None = Query(None),
):
    await require_project_access(project_id, current_user, db)
    return await test_plan_service.list_scenarios(db, project_id, requirement_id)


@router.get("/scenarios/{scenario_id}", response_model=TestScenarioOut)
async def get_scenario(
    scenario_id: int,
    db: DBSession,
    current_user: CurrentUser,
):
    sc = await test_plan_service.get_scenario(db, scenario_id)
    if not sc:
        raise HTTPException(status_code=404, detail="Scenario not found")
    await require_entity_project_access(sc, current_user, db)
    return sc


@router.post("/scenarios/{scenario_id}/approve", response_model=TestScenarioOut)
async def approve_scenario(
    scenario_id: int,
    body: ApprovalRequest,
    db: DBSession,
    current_user: CurrentUser,
):
    sc = await test_plan_service.get_scenario(db, scenario_id)
    if not sc:
        raise HTTPException(status_code=404, detail="Scenario not found")
    await require_permission(APPROVE_TEST_PLANS, sc.project_id, current_user, db)
    sc = await test_plan_service.approve_test_scenario(db, sc, body.action, body.notes)
    await approval_service.create_approval_action(
        db,
        project_id=sc.project_id,
        user_id=current_user.id,
        entity_type="test_scenario",
        entity_id=sc.id,
        action=body.action,
        notes=body.notes,
    )
    await db.commit()
    return sc


# ── Test Cases ────────────────────────────────────────────────────────────────

@router.get("/cases/project/{project_id}", response_model=list[TestCaseOut])
async def list_test_cases(
    project_id: int,
    db: DBSession,
    current_user: CurrentUser,
    scenario_id: int | None = Query(None),
    requirement_id: int | None = Query(None),
    status: str | None = Query(None),
    automation_only: bool = Query(False),
    skip: int = Query(0, ge=0),
    limit: int = Query(100, ge=1, le=500),
):
    await require_project_access(project_id, current_user, db)
    return await test_plan_service.list_test_cases(
        db, project_id, scenario_id, requirement_id, status, automation_only, skip, limit
    )


@router.get("/cases/project/{project_id}/summary", response_model=TestCaseSummaryOut)
async def test_case_project_summary(
    project_id: int,
    db: DBSession,
    current_user: CurrentUser,
):
    await require_project_access(project_id, current_user, db)
    return await test_plan_service.test_case_summary(db, project_id)


@router.get("/cases/{tc_id}", response_model=TestCaseOut)
async def get_test_case(
    tc_id: int,
    db: DBSession,
    current_user: CurrentUser,
):
    tc = await test_plan_service.get_test_case(db, tc_id)
    if not tc:
        raise HTTPException(status_code=404, detail="Test case not found")
    await require_entity_permission(tc, APPROVE_TEST_CASES, current_user, db)
    return tc


@router.patch("/cases/{tc_id}", response_model=TestCaseOut)
async def update_test_case(
    tc_id: int,
    updates: TestCaseUpdate,
    db: DBSession,
    current_user: CurrentUser,
):
    tc = await test_plan_service.get_test_case(db, tc_id)
    if not tc:
        raise HTTPException(status_code=404, detail="Test case not found")
    await require_entity_permission(tc, APPROVE_TEST_CASES, current_user, db)
    tc = await test_plan_service.update_test_case(db, tc, updates, user_id=current_user.id)
    await db.commit()
    return tc


@router.get("/cases/{tc_id}/history", response_model=list[TestCaseHistoryOut])
async def get_test_case_history(
    tc_id: int,
    db: DBSession,
    current_user: CurrentUser,
):
    tc = await test_plan_service.get_test_case(db, tc_id)
    if not tc:
        raise HTTPException(status_code=404, detail="Test case not found")
    await require_entity_project_access(tc, current_user, db)
    return await test_plan_service.list_test_case_history(db, tc)


@router.post("/cases/{tc_id}/sync-jira", response_model=TestCaseJiraSyncOut, status_code=202)
async def sync_test_case_jira(
    tc_id: int,
    db: DBSession,
    current_user: CurrentUser,
):
    tc = await test_plan_service.get_test_case(db, tc_id)
    if not tc:
        raise HTTPException(status_code=404, detail="Test case not found")
    await require_permission(SYNC_JIRA, tc.project_id, current_user, db)
    sync_job_id, task_id = await test_plan_service.sync_test_case_jira(db, tc, current_user.id)
    await db.commit()
    return TestCaseJiraSyncOut(test_case_id=tc.id, status="queued", sync_job_id=sync_job_id, task_id=task_id)


@router.post("/cases/{tc_id}/approve", response_model=TestCaseOut)
async def approve_test_case(
    tc_id: int,
    body: ApprovalRequest,
    db: DBSession,
    current_user: CurrentUser,
):
    tc = await test_plan_service.get_test_case(db, tc_id)
    if not tc:
        raise HTTPException(status_code=404, detail="Test case not found")
    await require_entity_permission(tc, APPROVE_TEST_CASES, current_user, db)
    tc = await test_plan_service.approve_test_case(db, tc, body.action, body.notes, current_user.id)
    await approval_service.create_approval_action(
        db,
        project_id=tc.project_id,
        user_id=current_user.id,
        entity_type="test_case",
        entity_id=tc.id,
        action=body.action,
        notes=body.notes,
    )
    await db.commit()
    return tc


# ── Agent Endpoints ───────────────────────────────────────────────────────────

@router.post("/agent/generate-plan")
async def trigger_planning_agent(
    body: AgentPlanTrigger,
    db: DBSession,
    current_user: CurrentUser,
):
    """
    Trigger Agent 3 (Test Planning) from approved requirements.
    Creates a structured test plan record.
    """
    await require_permission(APPROVE_TEST_PLANS, body.project_id, current_user, db)

    # Load project name
    result = await db.execute(select(Project).where(Project.id == body.project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Load requirements — only approved ones are allowed
    reqs = []
    not_approved = []
    wrong_project = []
    for rid in body.requirement_ids:
        result = await db.execute(select(Requirement).where(Requirement.id == rid))
        r = result.scalar_one_or_none()
        if not r:
            continue
        if r.project_id != body.project_id:
            wrong_project.append(rid)
            continue
        if r.status != "approved":
            not_approved.append(rid)
            continue
        reqs.append({
            "id": r.id,
            "requirement_id": r.requirement_id,
            "title": r.title,
            "summary": r.summary,
            "acceptance_criteria": r.acceptance_criteria,
            "business_rules": r.business_rules,
            "risks": r.risks,
        })

    if wrong_project:
        raise HTTPException(
            status_code=403,
            detail=f"Requirement ID(s) {wrong_project} do not belong to project {body.project_id}",
        )
    if not reqs:
        if not_approved:
            raise HTTPException(
                status_code=422,
                detail=f"Only approved requirements can be used for test planning. "
                       f"Not yet approved: {not_approved}",
            )
        raise HTTPException(status_code=422, detail="No valid requirements found")

    if not _run_agents_synchronously():
        agent_run, task_id = await enqueue_agent_run(
            db,
            project_id=body.project_id,
            user_id=current_user.id,
            agent_name="test_planning",
            input_data={"requirements": reqs, "project_name": project.name},
            metadata={"approved_requirement_ids": [r["id"] for r in reqs]},
        )
        return JSONResponse(
            status_code=202,
            content={"message": "Test planning queued", "agent_run_id": agent_run.id, "task_id": task_id},
        )

    user_id = current_user.id
    agent_run = await agent_run_service.start_agent_run(
        db,
        project_id=body.project_id,
        user_id=user_id,
        agent_name="test_planning",
        input_data=body.model_dump(mode="json"),
        metadata={"approved_requirement_ids": [r["id"] for r in reqs]},
    )

    agent = TestPlanningAgent()
    agent_result = await agent.run(requirements=reqs, project_name=project.name)

    if not agent_result.success:
        await agent_run_service.fail_agent_run(
            db,
            agent_run,
            error_message=agent_result.error or "Agent failed",
            agent_result=agent_result,
        )
        await db.commit()
        raise HTTPException(status_code=500, detail=agent_result.error or "Agent failed")

    plan_data = agent_result.data["test_plan"]

    plan = TestPlan(
        project_id=body.project_id,
        created_by=user_id,
        test_plan_id=temporary_id("TP"),
        title=plan_data.get("title", f"Test Plan — {project.name}"),
        scope=plan_data.get("scope"),
        out_of_scope=plan_data.get("out_of_scope"),
        test_types=plan_data.get("test_types"),
        entry_criteria=plan_data.get("entry_criteria"),
        exit_criteria=plan_data.get("exit_criteria"),
        risks=plan_data.get("risks"),
        mitigations=plan_data.get("mitigations"),
        automation_candidates=plan_data.get("automation_candidates"),
        estimated_effort=(
            plan_data["estimated_effort"]
            if isinstance(plan_data.get("estimated_effort"), str)
            else json.dumps(plan_data["estimated_effort"])
            if plan_data.get("estimated_effort") is not None
            else None
        ),
        resource_recommendation=(
            plan_data["resource_recommendation"]
            if isinstance(plan_data.get("resource_recommendation"), str)
            else json.dumps(plan_data["resource_recommendation"])
            if plan_data.get("resource_recommendation") is not None
            else None
        ),
        status="draft",
        agent_run_id=agent_run.id,
        metadata_={"source_requirement_ids": body.requirement_ids},
    )
    db.add(plan)
    await db.flush()
    plan.test_plan_id = display_id("TP", plan.id)
    await db.flush()
    await db.refresh(plan)
    await traceability_service.create_lineage_many(
        db,
        project_id=body.project_id,
        parents=[("requirement", r["id"]) for r in reqs],
        child_type="test_plan",
        child_id=plan.id,
        agent_run_id=agent_run.id,
    )
    await agent_run_service.complete_agent_run(
        db,
        agent_run,
        agent_result=agent_result,
        output_data={"plan_id": plan.id, "test_plan_id": plan.test_plan_id},
    )
    await db.commit()

    return {
        "message": "Test plan generated successfully",
        "plan_id": plan.id,
        "test_plan_id": plan.test_plan_id,
        "agent_logs": agent_result.logs,
        "agent_run_id": agent_run.id,
    }


@router.post("/agent/generate-scenarios")
async def trigger_scenario_agent(
    body: AgentPlanTrigger,
    db: DBSession,
    current_user: CurrentUser,
):
    """
    Trigger Agent 4 (Test Scenario) from requirements.
    Creates TestScenario records.
    """
    await require_permission(APPROVE_TEST_PLANS, body.project_id, current_user, db)

    reqs = []
    req_rows: list[Requirement] = []
    for rid in body.requirement_ids:
        result = await db.execute(select(Requirement).where(Requirement.id == rid))
        r = result.scalar_one_or_none()
        if r:
            if r.project_id != body.project_id:
                raise HTTPException(status_code=403, detail="Requirement does not belong to this project")
            req_rows.append(r)
            reqs.append({
                "id": r.id,
                "requirement_id": r.requirement_id,
                "title": r.title,
                "summary": r.summary,
                "acceptance_criteria": r.acceptance_criteria,
                "business_rules": r.business_rules,
                "user_roles": r.user_roles,
            })

    if not reqs:
        raise HTTPException(status_code=422, detail="No valid requirements found")

    # GAP-4c: block generation from failed/low-quality requirements unless overridden
    _enforce_quality_gate(req_rows, body.override_quality_gate)

    if not _run_agents_synchronously():
        agent_run, task_id = await enqueue_agent_run(
            db,
            project_id=body.project_id,
            user_id=current_user.id,
            agent_name="test_scenario",
            input_data={"requirements": reqs},
            metadata={"requirement_ids": [r["id"] for r in reqs]},
        )
        return JSONResponse(
            status_code=202,
            content={"message": "Test scenario generation queued", "agent_run_id": agent_run.id, "task_id": task_id},
        )

    user_id = current_user.id
    agent_run = await agent_run_service.start_agent_run(
        db,
        project_id=body.project_id,
        user_id=user_id,
        agent_name="test_scenario",
        input_data=body.model_dump(mode="json"),
        metadata={"requirement_ids": [r["id"] for r in reqs]},
    )

    agent = TestScenarioAgent()
    agent_result = await agent.run(
        requirements=reqs,
        db=db,
        project_id=body.project_id,
        agent_run_id=agent_run.id,
    )

    if not agent_result.success:
        await agent_run_service.fail_agent_run(
            db,
            agent_run,
            error_message=agent_result.error or "Agent failed",
            agent_result=agent_result,
        )
        await db.commit()
        raise HTTPException(status_code=500, detail=agent_result.error or "Agent failed")

    created_ids = []
    for i, sc_data in enumerate(agent_result.data.get("scenarios", [])):
        # Map back to DB requirement_id
        src_req_id = sc_data.get("_source_requirement_id")
        db_req_id = None
        if src_req_id:
            for r in reqs:
                if r["id"] == src_req_id or r["requirement_id"] == src_req_id:
                    db_req_id = r["id"]
                    break

        sc = TestScenario(
            project_id=body.project_id,
            requirement_id=db_req_id,
            created_by=user_id,
            scenario_id=temporary_id("TS"),
            title=sc_data.get("title", "Untitled Scenario"),
            description=sc_data.get("description"),
            scenario_type=sc_data.get("scenario_type", "positive"),
            priority=sc_data.get("priority", "Medium"),
            coverage_mapping=sc_data.get("coverage_mapping"),
            status="draft",
            agent_run_id=agent_run.id,
        )
        db.add(sc)
        await db.flush()
        sc.scenario_id = display_id("TS", sc.id)
        await db.flush()
        # Persist RAG citations if this scenario was grounded
        rag_chunk_ids = sc_data.get("_rag_chunk_ids")
        if rag_chunk_ids:
            try:
                from app.agents.rag_context import persist_citations
                from app.services.rag_retrieval_service import RetrievedChunk
                dummy_chunks = [
                    RetrievedChunk(
                        chunk_id=cid, chunk_text="", source_type="", source_id=None,
                        section=None, semantic_score=None, keyword_score=None, hybrid_score=0.0
                    )
                    for cid in rag_chunk_ids
                ]
                from app.agents.rag_context import RAGContext
                ctx = RAGContext(context_block="", citation_chunk_ids=rag_chunk_ids, grounded=True, chunks=dummy_chunks)
                await persist_citations(db, body.project_id, "test_scenario", sc.id, ctx, agent_run.id)
            except Exception:
                pass  # citations are best-effort — don't fail the request
        if db_req_id is not None:
            await traceability_service.create_lineage(
                db,
                project_id=body.project_id,
                parent_type="requirement",
                parent_id=db_req_id,
                child_type="test_scenario",
                child_id=sc.id,
                agent_run_id=agent_run.id,
            )
        created_ids.append(sc.id)

    await agent_run_service.complete_agent_run(
        db,
        agent_run,
        agent_result=agent_result,
        output_data={"scenario_ids": created_ids, "count": len(created_ids)},
    )
    await db.commit()
    return {
        "message": f"Generated {len(created_ids)} test scenarios",
        "scenario_ids": created_ids,
        "agent_logs": agent_result.logs,
        "agent_run_id": agent_run.id,
    }


@router.post("/agent/generate-cases")
async def trigger_test_case_agent(
    body: AgentCaseTrigger,
    db: DBSession,
    current_user: CurrentUser,
):
    """
    Trigger Agent 5 (Test Case Development) from test scenarios.
    Creates detailed TestCase records.
    """
    await require_permission(APPROVE_TEST_CASES, body.project_id, current_user, db)

    scenarios = []
    if body.scenario_ids:
        for sid in body.scenario_ids:
            result = await db.execute(select(TestScenario).where(TestScenario.id == sid))
            s = result.scalar_one_or_none()
            if s:
                if s.project_id != body.project_id:
                    raise HTTPException(status_code=403, detail="Scenario does not belong to this project")
                if s.status != "approved":
                    raise HTTPException(status_code=400, detail=f"Scenario {s.scenario_id} is not approved")
                scenarios.append({
                    "id": s.id,
                    "scenario_id": s.scenario_id,
                    "title": s.title,
                    "description": s.description,
                    "scenario_type": s.scenario_type,
                    "priority": s.priority,
                    "coverage_mapping": s.coverage_mapping,
                    "_source_requirement_id": s.requirement_id,
                })
    elif body.requirement_ids:
        for rid in body.requirement_ids:
            result = await db.execute(
                select(TestScenario).where(
                    TestScenario.requirement_id == rid,
                    TestScenario.status == "approved"
                )
            )
            for s in result.scalars().all():
                if s.project_id != body.project_id:
                    raise HTTPException(status_code=403, detail="Scenario does not belong to this project")
                scenarios.append({
                    "id": s.id,
                    "scenario_id": s.scenario_id,
                    "title": s.title,
                    "description": s.description,
                    "scenario_type": s.scenario_type,
                    "priority": s.priority,
                    "coverage_mapping": s.coverage_mapping,
                    "_source_requirement_id": s.requirement_id,
                })

    # Fallback: if requirement-based lookup yielded nothing, use all project scenarios
    if not scenarios and body.requirement_ids:
        result = await db.execute(
            select(TestScenario).where(
                TestScenario.project_id == body.project_id,
                TestScenario.status == "approved"
            )
        )
        for s in result.scalars().all():
            scenarios.append({
                "id": s.id,
                "scenario_id": s.scenario_id,
                "title": s.title,
                "description": s.description,
                "scenario_type": s.scenario_type,
                "priority": s.priority,
                "coverage_mapping": s.coverage_mapping,
                "_source_requirement_id": s.requirement_id,
            })

    if not scenarios:
        raise HTTPException(status_code=422, detail="No scenarios found to generate test cases from")

    # GAP-4c: gate on the parent requirements of the selected scenarios
    gate_req_ids = {s.get("_source_requirement_id") for s in scenarios if s.get("_source_requirement_id")}
    if body.requirement_ids:
        gate_req_ids.update(body.requirement_ids)
    if gate_req_ids:
        gate_rows_result = await db.execute(
            select(Requirement).where(
                Requirement.id.in_(gate_req_ids),
                Requirement.project_id == body.project_id,
            )
        )
        _enforce_quality_gate(list(gate_rows_result.scalars().all()), body.override_quality_gate)

    if not _run_agents_synchronously():
        agent_run, task_id = await enqueue_agent_run(
            db,
            project_id=body.project_id,
            user_id=current_user.id,
            agent_name="test_case",
            input_data={"scenarios": scenarios},
            metadata={"scenario_ids": [s["id"] for s in scenarios]},
        )
        return JSONResponse(
            status_code=202,
            content={"message": "Test case generation queued", "agent_run_id": agent_run.id, "task_id": task_id},
        )

    user_id = current_user.id
    agent_run = await agent_run_service.start_agent_run(
        db,
        project_id=body.project_id,
        user_id=user_id,
        agent_name="test_case",
        input_data=body.model_dump(mode="json"),
        metadata={"scenario_ids": [s["id"] for s in scenarios]},
    )

    agent = TestCaseDevelopmentAgent()
    agent_result = await agent.run(scenarios=scenarios)

    if not agent_result.success:
        await agent_run_service.fail_agent_run(
            db,
            agent_run,
            error_message=agent_result.error or "Agent failed",
            agent_result=agent_result,
        )
        await db.commit()
        raise HTTPException(status_code=500, detail=agent_result.error or "Agent failed")

    # Count existing test cases for IDs
    # Build scenario ID → DB id map
    scenario_id_map = {s["scenario_id"]: s["id"] for s in scenarios}
    req_id_map = {s["id"]: s["_source_requirement_id"] for s in scenarios}

    created_ids = []
    for i, tc_data in enumerate(agent_result.data.get("test_cases", [])):
        src_sc_id = tc_data.get("_source_scenario_id")
        db_sc_id = scenario_id_map.get(src_sc_id) if src_sc_id else None
        db_req_id = req_id_map.get(db_sc_id) if db_sc_id else None

        tc = TestCase(
            project_id=body.project_id,
            scenario_id=db_sc_id,
            requirement_id=db_req_id,
            created_by=user_id,
            test_case_id=temporary_id("TC"),
            title=tc_data.get("title", "Untitled Test Case"),
            preconditions=tc_data.get("preconditions"),
            test_data=tc_data.get("test_data"),
            steps=tc_data.get("steps"),
            expected_result=tc_data.get("expected_result"),
            bdd_scenario=tc_data.get("bdd_scenario"),
            priority=tc_data.get("priority", "Medium"),
            severity=tc_data.get("severity", "Medium"),
            test_type=tc_data.get("test_type"),
            automation_candidate=bool(tc_data.get("automation_candidate", False)),
            metadata_={"tags": tc_data.get("tags")} if tc_data.get("tags") else None,
            status="draft",
            agent_run_id=agent_run.id,
        )
        db.add(tc)
        await db.flush()
        tc.test_case_id = display_id("TC", tc.id)
        await db.flush()
        parents = []
        if db_sc_id is not None:
            parents.append(("test_scenario", db_sc_id))
        if db_req_id is not None:
            parents.append(("requirement", db_req_id))
        await traceability_service.create_lineage_many(
            db,
            project_id=body.project_id,
            parents=parents,
            child_type="test_case",
            child_id=tc.id,
            agent_run_id=agent_run.id,
        )
        created_ids.append(tc.id)

    await agent_run_service.complete_agent_run(
        db,
        agent_run,
        agent_result=agent_result,
        output_data={"test_case_ids": created_ids, "count": len(created_ids)},
    )
    await db.commit()
    return {
        "message": f"Generated {len(created_ids)} test cases",
        "test_case_ids": created_ids,
        "agent_logs": agent_result.logs,
        "agent_run_id": agent_run.id,
    }


@router.get("/{plan_id}/export/docx")
async def export_test_plan_docx(
    plan_id: int,
    db: DBSession,
    current_user: CurrentUser,
):
    plan = await test_plan_service.get_test_plan(db, plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="Test plan not found")
    await require_entity_permission(plan, APPROVE_TEST_PLANS, current_user, db)

    # Generate DOCX
    doc = Document()
    doc.add_heading(plan.title, 0)

    doc.add_heading("Metadata", level=1)
    doc.add_paragraph(f"Test Plan ID: {plan.test_plan_id}")
    doc.add_paragraph(f"Status: {plan.status}")
    doc.add_paragraph(f"Estimated Effort: {plan.estimated_effort or 'N/A'}")
    doc.add_paragraph(f"Resource Recommendation: {plan.resource_recommendation or 'N/A'}")

    sections = [
        ("Scope", plan.scope),
        ("Out of Scope", plan.out_of_scope),
        ("Test Types", plan.test_types),
        ("Entry Criteria", plan.entry_criteria),
        ("Exit Criteria", plan.exit_criteria),
        ("Risks", plan.risks),
        ("Mitigations", plan.mitigations),
        ("Automation Candidates", plan.automation_candidates),
    ]

    for title, items in sections:
        doc.add_heading(title, level=1)
        if items:
            for item in items:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph("N/A")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=test_plan_{plan.test_plan_id}.docx"}
    )

