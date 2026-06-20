"""Document storage and extraction services."""
from __future__ import annotations

import os
import re
import uuid
from pathlib import Path

import aiofiles
from fastapi import HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models.document import UploadedDocument
from app.models.project import Project

settings = get_settings()

ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/markdown": "md",
    "text/csv": "csv",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    # GAP-1: UI screenshot/image inputs
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/webp": "webp",
}

EXT_MAP = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".txt": "txt",
    ".md": "md",
    ".csv": "csv",
    ".xlsx": "xlsx",
    # GAP-1: UI screenshot/image inputs
    ".png": "png",
    ".jpg": "jpg",
    ".jpeg": "jpg",
    ".webp": "webp",
}

# GAP-1: file types treated as UI screenshots (analysed by vision agent,
# not by text extraction)
IMAGE_TYPES = {"png", "jpg", "webp"}

MAX_SIZE = settings.max_upload_size_mb * 1024 * 1024
CHUNK_SIZE = 1024 * 1024
SAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9._-]+")


def sanitize_filename(filename: str | None) -> str:
    raw_name = Path(filename or "upload").name
    safe_name = SAFE_FILENAME_RE.sub("_", raw_name).strip("._")
    return safe_name or "upload"


def validate_file_signature(file_type: str, first_bytes: bytes) -> None:
    if file_type == "pdf" and not first_bytes.startswith(b"%PDF-"):
        raise HTTPException(status_code=415, detail="File content is not a valid PDF")
    # GAP-1: image magic-byte validation
    if file_type == "png" and not first_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
        raise HTTPException(status_code=415, detail="File content is not a valid PNG image")
    if file_type == "jpg" and not first_bytes.startswith(b"\xff\xd8\xff"):
        raise HTTPException(status_code=415, detail="File content is not a valid JPEG image")
    if file_type == "webp" and not (first_bytes[:4] == b"RIFF" and first_bytes[8:12] == b"WEBP"):
        raise HTTPException(status_code=415, detail="File content is not a valid WebP image")
    if file_type in {"docx", "xlsx"} and not first_bytes.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
        raise HTTPException(status_code=415, detail=f"File content is not a valid {file_type.upper()} archive")
    if file_type in {"txt", "md", "csv"}:
        if b"\x00" in first_bytes:
            raise HTTPException(status_code=415, detail="Text file appears to contain binary data")
        try:
            first_bytes.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise HTTPException(status_code=415, detail="Text file is not valid UTF-8") from exc


async def upload_document(
    db: AsyncSession,
    project_id: int,
    user_id: int,
    file: UploadFile,
) -> UploadedDocument:
    """Stream a document to disk, create a DB row, and enqueue extraction."""
    safe_original_name = sanitize_filename(file.filename)
    ext = Path(safe_original_name).suffix.lower()
    file_type = ALLOWED_TYPES.get(file.content_type or "") or EXT_MAP.get(ext)
    if not file_type:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type: {file.content_type} / {ext}",
        )

    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    storage_dir = Path(settings.file_storage_path) / "uploads" / str(project_id)
    storage_dir.mkdir(parents=True, exist_ok=True)

    stored_name = f"{uuid.uuid4().hex}_{safe_original_name}"
    file_path = storage_dir / stored_name

    size = 0
    first_chunk = await file.read(CHUNK_SIZE)
    if not first_chunk:
        raise HTTPException(status_code=422, detail="Uploaded file is empty")
    validate_file_signature(file_type, first_chunk[:8192])

    async with aiofiles.open(file_path, "wb") as f:
        chunk = first_chunk
        while chunk:
            size += len(chunk)
            if size > MAX_SIZE:
                try:
                    os.remove(file_path)
                except OSError:
                    pass
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=f"File exceeds {settings.max_upload_size_mb} MB limit",
                )
            await f.write(chunk)
            chunk = await file.read(CHUNK_SIZE)

    # GAP-1: images skip text extraction — they are analysed on demand by the
    # UI analysis (vision) agent. Mark them processed immediately.
    is_image = file_type in IMAGE_TYPES

    doc = UploadedDocument(
        project_id=project_id,
        created_by=user_id,
        original_filename=safe_original_name,
        stored_filename=stored_name,
        file_path=str(file_path),
        file_type=file_type,
        file_size_bytes=size,
        status="processed" if is_image else "processing",
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)
    await db.commit()
    await db.refresh(doc)

    if is_image:
        doc.metadata_ = {"input_kind": "ui_image"}
        await db.commit()
        await db.refresh(doc)
        return doc

    from app.worker.tasks.document_tasks import extract_document_text

    task = extract_document_text.delay(doc.id)
    doc.metadata_ = {"extraction_task_id": task.id}
    await db.commit()
    await db.refresh(doc)
    return doc


def sanitize_extracted_text(text: str) -> str:
    """Neutralize potential prompt injections and truncate length for LLM safety."""
    if not text:
        return ""
    
    # 1. Truncate to limit length passed to LLM (e.g. 50,000 characters)
    max_len = 50000
    text = text[:max_len]
    
    # 2. Strip potential prompt injection patterns
    from app.security.prompt_guard import INJECTION_PATTERNS
    import re
    for pattern in INJECTION_PATTERNS:
        text = re.sub(pattern, "[SANITIZED PAYLOAD]", text)
        
    return text


def extract_text(file_path: str, file_type: str, contents: bytes | None = None) -> tuple[str, int | None]:
    """Extract plain text from a stored document."""
    if contents is None:
        contents = Path(file_path).read_bytes()

    text = ""
    pages = None

    if file_type == "pdf":
        text, pages = _extract_pdf(contents)
    elif file_type == "docx":
        text, pages = _extract_docx(contents)
    elif file_type in ("txt", "md", "csv"):
        text, pages = contents.decode("utf-8", errors="replace"), None
    elif file_type == "xlsx":
        text, pages = _extract_xlsx(contents)

    # Sanitize and truncate the extracted text
    text = sanitize_extracted_text(text)

    return text, pages



def _extract_pdf(contents: bytes) -> tuple[str, int]:
    import fitz

    doc = fitz.open(stream=contents, filetype="pdf")
    pages = [page.get_text() for page in doc]
    return "\n\n".join(pages), len(doc)


def _extract_docx(contents: bytes) -> tuple[str, None]:
    import io

    from docx import Document

    doc = Document(io.BytesIO(contents))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs), None


def _extract_xlsx(contents: bytes) -> tuple[str, None]:
    import io

    import pandas as pd

    xls = pd.ExcelFile(io.BytesIO(contents))
    sheets = []
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        sheets.append(f"## Sheet: {sheet}\n{df.to_string(index=False)}")
    return "\n\n".join(sheets), None


async def delete_document(db: AsyncSession, doc_id: int) -> None:
    result = await db.execute(select(UploadedDocument).where(UploadedDocument.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    try:
        if doc.file_path and os.path.exists(doc.file_path):
            os.remove(doc.file_path)
    except OSError:
        pass
    await db.delete(doc)
    await db.flush()


async def get_document(db: AsyncSession, doc_id: int, project_id: int | None = None) -> UploadedDocument | None:
    stmt = select(UploadedDocument).where(UploadedDocument.id == doc_id)
    if project_id:
        stmt = stmt.where(UploadedDocument.project_id == project_id)
    result = await db.execute(stmt)
    return result.scalar_one_or_none()


async def list_documents(
    db: AsyncSession,
    project_id: int,
    skip: int = 0,
    limit: int = 100,
) -> list[UploadedDocument]:
    stmt = (
        select(UploadedDocument)
        .where(UploadedDocument.project_id == project_id)
        .order_by(UploadedDocument.created_at.desc())
        .offset(skip)
        .limit(limit)
    )
    result = await db.execute(stmt)
    return list(result.scalars().all())
